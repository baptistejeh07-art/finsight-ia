# -*- coding: utf-8 -*-
"""One-pager Word FinSight IA — version visuelle adaptée jury L2 éco-gestion."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
NAVY_HEX = '1B3A6B'
INK = RGBColor(0x1F, 0x29, 0x37)
INK_500 = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT = RGBColor(0xB0, 0x60, 0x00)

doc = Document()

for section in doc.sections:
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = INK


def shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_border_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0.1, bottom=0.1, left=0.2, right=0.2):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(int(val * 567)))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


# =========================================================
# HEADER — bande navy pleine largeur
# =========================================================
header_tbl = doc.add_table(rows=1, cols=1)
header_tbl.autofit = False
header_tbl.columns[0].width = Cm(18)
cell = header_tbl.rows[0].cells[0]
shade(cell, NAVY_HEX)
set_cell_margins(cell, top=0.35, bottom=0.35, left=0.5, right=0.5)

p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
r1 = p.add_run("FinSight IA")
r1.font.size = Pt(26)
r1.font.bold = True
r1.font.color.rgb = WHITE
r2 = p.add_run("    Plateforme d\u2019analyse financi\u00e8re augment\u00e9e par IA")
r2.font.size = Pt(12)
r2.font.italic = True
r2.font.color.rgb = RGBColor(0xC5, 0xD5, 0xE8)

# Ligne sous-titre : projet porté par + URL à droite
sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(4)
sub.paragraph_format.space_after = Pt(6)
r = sub.add_run("Projet entrepreneurial port\u00e9 par Baptiste Jehanno \u2014 BTS CG en alternance (Toulouse)      ")
r.font.size = Pt(10)
r.font.color.rgb = INK_500
r2 = sub.add_run("finsight-ia.com")
r2.font.size = Pt(10)
r2.font.bold = True
r2.font.color.rgb = NAVY

# =========================================================
# CHIFFRES CLÉS — 4 tuiles
# =========================================================
tiles_tbl = doc.add_table(rows=1, cols=4)
tiles_tbl.autofit = False
for i in range(4):
    tiles_tbl.columns[i].width = Cm(4.5)

metrics = [
    ("< 60 sec", "pour g\u00e9n\u00e9rer un rapport complet"),
    ("25 graphiques", "modulables en glisser-d\u00e9poser"),
    ("+8,9 %", "alpha statistiquement significatif"),
    ("6\u00a0500", "observations historiques test\u00e9es"),
]
for i, (val, label) in enumerate(metrics):
    c = tiles_tbl.rows[0].cells[i]
    shade(c, 'F3F6FA')
    set_cell_border_none(c)
    set_cell_margins(c, top=0.25, bottom=0.25, left=0.25, right=0.25)

    c.text = ""
    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run(val)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(label)
    r.font.size = Pt(8.5)
    r.font.color.rgb = INK_500


def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'D1D5DB')
    pbdr.append(bottom)
    pPr.append(pbdr)


# =========================================================
# LE PROJET EN CHIFFRES
# =========================================================
section_title("Le projet en une ligne")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.2
r = p.add_run(
    "D\u00e9mocratiser l\u2019analyse financi\u00e8re fondamentale\u00a0: un outil qui produit, en moins d\u2019une minute, "
    "le m\u00eame livrable qu\u2019un analyste junior en deux \u00e0 trois heures \u2014\u00a0pour moins de 35\u00a0\u20ac par mois."
)
r.font.size = Pt(10.5)

# =========================================================
# COMMENT ÇA MARCHE — 3 étapes visuelles
# =========================================================
section_title("Comment \u00e7a marche")

flow_tbl = doc.add_table(rows=2, cols=3)
flow_tbl.autofit = False
for i in range(3):
    flow_tbl.columns[i].width = Cm(6)

steps = [
    ("1.", "L\u2019utilisateur saisit un ticker (AAPL), un secteur, un indice ou un num\u00e9ro SIREN pour une PME."),
    ("2.", "Un pipeline de sept agents IA r\u00e9cup\u00e8re les donn\u00e9es, calcule les ratios, produit une synth\u00e8se et la stress-teste."),
    ("3.", "L\u2019utilisateur t\u00e9l\u00e9charge un rapport PDF, un pitchbook PowerPoint, un mod\u00e8le Excel, et explore un tableau de bord interactif."),
]
for i, (num, text) in enumerate(steps):
    # ligne 1 : numéro dans cercle/bloc navy
    top = flow_tbl.rows[0].cells[i]
    shade(top, '1B3A6B')
    set_cell_border_none(top)
    set_cell_margins(top, top=0.15, bottom=0.15, left=0.2, right=0.2)
    top.text = ""
    pt = top.paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(0)
    pt.paragraph_format.space_after = Pt(0)
    rt = pt.add_run(f"\u00c9TAPE {num.rstrip('.')}")
    rt.font.size = Pt(10)
    rt.font.bold = True
    rt.font.color.rgb = WHITE

    # ligne 2 : description
    bot = flow_tbl.rows[1].cells[i]
    shade(bot, 'F3F6FA')
    set_cell_border_none(bot)
    set_cell_margins(bot, top=0.2, bottom=0.25, left=0.25, right=0.25)
    bot.text = ""
    pb = bot.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(0)
    pb.paragraph_format.line_spacing = 1.15
    rb = pb.add_run(text)
    rb.font.size = Pt(9.5)

# =========================================================
# RÉSULTATS VÉRIFIÉS — Score FinSight
# =========================================================
section_title("R\u00e9sultats v\u00e9rifi\u00e9s \u2014 le score propri\u00e9taire test\u00e9 sur 10 ans")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(3)
r = p.add_run(
    "J\u2019ai construit un score propri\u00e9taire qui attribue une note 0\u2013100 \u00e0 chaque soci\u00e9t\u00e9 selon quatre "
    "dimensions (qualit\u00e9, valeur, momentum, risque). Test\u00e9 historiquement sur les 100 plus grandes "
    "capitalisations am\u00e9ricaines sur dix ans, il produit\u00a0:"
)
r.font.size = Pt(10)

results = doc.add_table(rows=4, cols=3)
results.autofit = False
widths = [Cm(6.5), Cm(4.5), Cm(7.0)]
for i, w in enumerate(widths):
    results.columns[i].width = w

hdr = results.rows[0].cells
for i, text in enumerate(["Profil testé", "Surperformance annuelle", "Fiabilité statistique"]):
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_margins(hdr[i], top=0.15, bottom=0.15, left=0.2, right=0.2)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    shade(hdr[i], NAVY_HEX)
    set_cell_border_none(hdr[i])

rows_data = [
    ("Profil \u00e9quilibr\u00e9", "+8,9\u00a0%", "Seuil de 95\u00a0% atteint (t = +2,10)"),
    ("Croissance (Tech)", "+19,4\u00a0%", "\u00c9chantillon restreint"),
    ("Valeur (cycliques)", "+24,0\u00a0%", "Conforme \u00e0 la litt\u00e9rature acad\u00e9mique"),
]
for row_idx, row_data in enumerate(rows_data, start=1):
    cells = results.rows[row_idx].cells
    for i, val in enumerate(row_data):
        set_cell_border_none(cells[i])
        set_cell_margins(cells[i], top=0.12, bottom=0.12, left=0.2, right=0.2)
        bg = 'FFFFFF' if row_idx % 2 == 1 else 'F9FAFB'
        shade(cells[i], bg)
        p = cells[i].paragraphs[0]
        if i == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(9.5)
        if i == 0:
            r.font.bold = True
        if i == 1:
            r.font.bold = True
            r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.15
r = p.add_run(
    "Limites pleinement assum\u00e9es\u00a0: la fen\u00eatre 2015-2025 favorise les valeurs de croissance. "
    "L\u2019extension \u00e0 la p\u00e9riode 2000-2025 est pr\u00e9vue pour couvrir \u00e9galement la crise de 2008 et la rotation value."
)
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = INK_500

# Lien méthodologie complète — compact
note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(3)
note_p.paragraph_format.space_after = Pt(0)
note_p.paragraph_format.line_spacing = 1.1
note_r1 = note_p.add_run("M\u00e9thodologie compl\u00e8te ")
note_r1.font.size = Pt(9)
note_r1.font.bold = True
note_r1.font.color.rgb = NAVY
note_r2 = note_p.add_run("(formules, protocole backtest, gouvernance IA, choix de conception)\u00a0: ")
note_r2.font.size = Pt(9)
note_r2.font.italic = True
note_r2.font.color.rgb = INK_500
note_r3 = note_p.add_run("finsight-ia.com/methodologie")
note_r3.font.size = Pt(9)
note_r3.font.bold = True
note_r3.font.color.rgb = NAVY
note_r3.font.underline = True

# =========================================================
# CE QUE CE PROJET M'A APPRIS — grille 2x3
# =========================================================
section_title("Ce que ce projet m\u2019a appris")

comp_tbl = doc.add_table(rows=3, cols=2)
comp_tbl.autofit = False
comp_tbl.columns[0].width = Cm(9)
comp_tbl.columns[1].width = Cm(9)

competences = [
    (
        "Finance d\u2019entreprise",
        "Valorisation DCF, ratios sectoriels, solvabilit\u00e9 (Altman\u00a0Z), analyse comparative de peers."
    ),
    (
        "M\u00e9thode scientifique",
        "Tests de significativit\u00e9 (t-stat, p-value), d\u00e9tection et correction des biais m\u00e9thodologiques."
    ),
    (
        "Gestion de projet",
        "Du cahier des charges \u00e0 la production continue, arbitrages techniques et prise de d\u00e9cision autonome."
    ),
    (
        "Connaissance du march\u00e9",
        "\u00c9tude concurrentielle (Bloomberg, Factset, Koyfin), tarification SaaS, segmentation client."
    ),
    (
        "Rigueur intellectuelle",
        "Argumentation chiffr\u00e9e, limites assum\u00e9es, absence d\u2019affirmation commerciale non v\u00e9rifi\u00e9e."
    ),
    (
        "Autonomie technique",
        "Construction int\u00e9grale du produit en dix-huit mois, d\u00e9ploiement en production et support utilisateur."
    ),
]

for i, (titre, desc) in enumerate(competences):
    row_idx = i // 2
    col_idx = i % 2
    cell = comp_tbl.rows[row_idx].cells[col_idx]
    shade(cell, 'F8F9FB')
    set_cell_border_none(cell)
    set_cell_margins(cell, top=0.2, bottom=0.2, left=0.25, right=0.25)
    cell.text = ""

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(titre)
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(desc)
    r2.font.size = Pt(9.5)

# =========================================================
# FOOTER
# =========================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '10')
top.set(qn('w:color'), '1B3A6B')
pbdr.append(top)
pPr.append(pbdr)
r = p.add_run("\nfinsight-ia.com")
r.font.bold = True
r.font.size = Pt(11)
r.font.color.rgb = NAVY
r2 = p.add_run("      \u00b7      baptiste.jeh07@gmail.com      \u00b7      07 49 28 48 81")
r2.font.size = Pt(9.5)
r2.font.color.rgb = INK_500

out_path = r'C:\Users\bapti\OneDrive\Perso\FinSight_One_Pager_Candidature_L2_v5.docx'
doc.save(out_path)
import os
print(f'OK saved: {out_path} ({os.path.getsize(out_path) // 1024} Ko)')
